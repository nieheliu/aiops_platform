# -*- coding: utf-8 -*-
from docx import Document

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
doc = Document(r"d:\javawebworkspace\aiops\组员B-课程设计报告-已填入.docx")

for i, p in enumerate(doc.paragraphs):
    if i < 35 or i > 250:
        continue
    t = p.text.strip()
    has_img = any(el.tag.endswith("}drawing") for el in p._p.iter())
    if t or has_img:
        suffix = " [含图片]" if has_img else ""
        print(f"{i}{suffix}|{t}")

print("\n--- 表格标题段落 ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith("表3-"):
        print(i, t)
