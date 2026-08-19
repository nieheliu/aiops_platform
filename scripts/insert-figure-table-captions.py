# -*- coding: utf-8 -*-
"""在 Word 报告中图题/表题下方插入说明段落（生成新文件，不覆盖原稿）。"""
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "组员B-课程设计报告-已填入.docx"
OUT = ROOT / "组员B-课程设计报告-含图表说明.docx"

# 图题关键字 -> 说明（图题段落需包含关键字）
FIGURE_CAPTIONS = {
    "系统整体逻辑架构图": (
        "如图1所示，系统整体逻辑架构自上而下分为接入层、业务层、数据访问层与基础设施层。"
        "外部 Alertmanager 通过 Webhook 将告警推送至后端；业务层完成告警、工单、AI 诊断与知识检索；"
        "MySQL 持久化核心数据，RabbitMQ 异步建单，Redis 缓存统计，Elasticsearch 支撑知识检索，"
        "Spring AI 调用大模型 API，体现监控驱动业务闭环的设计思想。"
    ),
    "数据库E-R": (
        "如图14所示，数据库共 9 张核心表：用户与角色多对多关联，告警与工单一对一路由，"
        "工单关联操作日志与 AI 诊断，知识库可追溯告警/工单/诊断来源，外键保证业务链路可审计。"
    ),
}

# 按图片出现顺序（告警接入起共 12 张模块图，在 E-R 之前）
MODULE_FIGURE_TEXTS = [
    # 告警接入 活动、类
    "如图2所示，告警接入流程为：接收 firing 告警 → 解析入库 → 发送 MQ 消息 → 快速返回，仅处理触发状态告警。",
    "如图3所示，告警接入核心类包括 AlertWebhookController、AlertWorkflowServiceImpl、AlertTicketProducer 与 OpsAlert 实体。",
    # 异步建单
    "如图4所示，消费者按 alertId 幂等建单，成功 manualAck，失败进入死信队列，避免 Webhook 同步阻塞。",
    "如图5所示，异步建单涉及 RabbitMqConfig、Producer、Consumer 与 AlertWorkflowServiceImpl 协作完成事件驱动建单。",
    # 工单流转
    "如图6所示，工单按 0 待处理→1 处理中→2 已解决→3 已关闭 流转，每次操作写入 ops_ticket_log 审计。",
    "如图7所示，OpsTicketController 与 TicketKnowledgeServiceImpl 封装状态机及处理人归属权限校验。",
    # AI 诊断
    "如图8所示，AI 诊断收集告警与工单上下文构建 Prompt，调用大模型后将根因与修复建议写入 ai_diagnosis 表。",
    "如图9所示，AiDiagnosisWorkflowServiceImpl 通过 AiChatClientFactory 支持多模型接入与同工单同模型幂等。",
    # 知识库
    "如图10所示，工单解决时知识写入 MySQL 并同步 ES，前端关键词检索支持多字段高亮展示。",
    "如图11所示，TicketKnowledgeServiceImpl 与 KnowledgeArticleServiceImpl 实现双存储与生命周期管理。",
    # 认证授权
    "如图12所示，登录签发 JWT，拦截器校验 Token，Service 层叠加 RBAC 与业务归属双重鉴权。",
    "如图13所示，AuthController、JwtAuthInterceptor、JwtTokenProvider 与 SysUserServiceImpl 构成认证授权体系。",
]

TABLE_CAPTIONS = {
    "表3-1 告警接入模块核心类说明": (
        "如表3-1所示，告警接入由 Controller 接收入口、Service 编排入库、MQ 发送消息、Entity 映射 ops_alert 表，职责分层清晰。"
    ),
    "表3-2 工单流转模块接口说明": (
        "如表3-2所示，start/assign/resolve/close 四个接口分别对应认领、改派、解决与关闭，驱动四态工单状态机。"
    ),
    "表3-1 sys_user": (
        "如表3-1所示，sys_user 存储登录账号与状态，username 唯一，密码 BCrypt 加密，是权限与处理人体系的基础表。"
    ),
    "表3-2 sys_role": (
        "如表3-2所示，sys_role 定义 ADMIN/OPS/VIEWER 三种角色，role_code 作为程序内权限判断依据。"
    ),
    "表3-3 sys_user_role": (
        "如表3-3所示，用户与角色多对多关联表，本项目约束为单用户单角色，外键级联维护一致性。"
    ),
    "表3-4 ops_alert": (
        "如表3-4所示，ops_alert 记录接入告警，severity 表等级，raw_payload 保留 JSON 原始上下文供 AI 使用。"
    ),
    "表3-5 ops_ticket": (
        "如表3-5所示，ops_ticket 通过 alert_id 关联告警、handler_user_id 绑定处理人，status 驱动四态状态机。"
    ),
    "表3-6 ops_ticket_log": (
        "如表3-6所示，ops_ticket_log 记录 CREATE/START/RESOLVE/CLOSE 等操作，实现工单全流程审计。"
    ),
    "表3-7 ai_diagnosis": (
        "如表3-7所示，ai_diagnosis 保存大模型输出的根因分析与修复建议，关联告警与工单。"
    ),
    "表3-8 ops_knowledge": (
        "如表3-8所示，ops_knowledge 以 Markdown 存储知识正文，记录来源与 ES 同步及生命周期状态。"
    ),
    "表3-9 ops_knowledge_audit_log": (
        "如表3-9所示，知识审计表记录创建、审核、发布、归档等操作及状态变更，保证知识库可追溯。"
    ),
}


def insert_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    run = new_para.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体"
    )
    return new_para


def has_image(paragraph):
    return any(el.tag.endswith("}drawing") for el in paragraph._p.iter())


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))

    module_fig_idx = 0
    for p in doc.paragraphs:
        t = p.text.strip()

        for key, desc in FIGURE_CAPTIONS.items():
            if key in t and "如图" not in t:
                insert_after(p, desc)
                break

        for key, desc in TABLE_CAPTIONS.items():
            if t.startswith(key) or key in t:
                if "如表" not in t and len(t) < 80:
                    insert_after(p, desc)
                break

        if has_image(p):
            if module_fig_idx < len(MODULE_FIGURE_TEXTS):
                insert_after(p, MODULE_FIGURE_TEXTS[module_fig_idx])
                module_fig_idx += 1

    doc.save(str(OUT))
    print(f"已生成: {OUT}")
    print(f"模块图说明插入: {module_fig_idx} 处")


if __name__ == "__main__":
    main()
