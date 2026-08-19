# -*- coding: utf-8 -*-
"""将组员 B 的报告内容（3.1 模块、3.3 数据表、第6章总结）写入 Word 报告。"""
import copy
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "《软件工程课程设计III 》课程设计报告-模板.docx"
OUTPUT = ROOT / "组员B-课程设计报告-已填入.docx"


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体"
        )
    if style:
        new_para.style = style
    return new_para


def set_paragraph_text(paragraph, text):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体"
    )


def add_table_after(doc, paragraph, caption, headers, rows):
    """在段落后插入带标题的数据表。"""
    cap = insert_paragraph_after(paragraph, caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.bold = True
        run.font.size = Pt(10.5)

    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10.5)

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            tbl.rows[ri + 1].cells[ci].text = str(val)
            for p in tbl.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10.5)

  # 将表格移到 caption 后面
    paragraph._p.addnext(tbl._tbl)
    return tbl


def find_paragraph(doc, keyword):
    for p in doc.paragraphs:
        if keyword in p.text:
            return p
    return None


def find_paragraph_in_range(doc, keyword, start_idx=0):
    for i, p in enumerate(doc.paragraphs):
        if i >= start_idx and keyword in p.text:
            return p, i
    return None, -1


# ---------- 3.3 数据表定义 ----------
DB_TABLES = [
    (
        "表3-1 sys_user（用户信息表）",
        ["字段名", "数据类型", "约束", "说明"],
        [
            ["id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "用户主键"],
            ["username", "VARCHAR(64)", "NOT NULL, UNIQUE", "登录用户名"],
            ["password", "VARCHAR(255)", "NOT NULL", "密码（BCrypt加密）"],
            ["email", "VARCHAR(128)", "UNIQUE", "邮箱"],
            ["status", "TINYINT", "NOT NULL, DEFAULT 1", "0-禁用，1-启用"],
            ["create_time", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "创建时间"],
            ["update_time", "DATETIME", "ON UPDATE CURRENT_TIMESTAMP", "更新时间"],
        ],
    ),
    (
        "表3-2 sys_role（角色定义表）",
        ["字段名", "数据类型", "约束", "说明"],
        [
            ["id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "角色主键"],
            ["role_name", "VARCHAR(64)", "NOT NULL, UNIQUE", "角色名称"],
            ["role_code", "VARCHAR(64)", "NOT NULL, UNIQUE", "ADMIN/OPS/VIEWER"],
            ["description", "VARCHAR(255)", "", "角色描述"],
            ["create_time", "DATETIME", "", "创建时间"],
            ["update_time", "DATETIME", "", "更新时间"],
        ],
    ),
    (
        "表3-3 sys_user_role（用户角色关联表）",
        ["字段名", "数据类型", "约束", "说明"],
        [
            ["user_id", "BIGINT UNSIGNED", "PK, FK→sys_user.id", "用户ID"],
            ["role_id", "BIGINT UNSIGNED", "PK, FK→sys_role.id", "角色ID"],
            ["grant_time", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "授权时间"],
        ],
    ),
    (
        "表3-4 ops_alert（告警事件表）",
        ["字段名", "数据类型", "约束", "说明"],
        [
            ["id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "告警主键"],
            ["alert_name", "VARCHAR(128)", "NOT NULL", "告警名称"],
            ["severity", "TINYINT", "NOT NULL", "1低/2中/3高/4紧急"],
            ["instance_ip", "VARCHAR(45)", "NOT NULL", "告警实例IP"],
            ["raw_payload", "JSON", "", "Alertmanager原始载荷"],
            ["trigger_time", "DATETIME", "NOT NULL", "触发时间"],
            ["create_time", "DATETIME", "", "入库时间"],
        ],
    ),
    (
        "表3-5 ops_ticket（运维工单表）",
        ["字段名", "数据类型", "约束", "说明"],
        [
            ["id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "工单主键"],
            ["alert_id", "BIGINT UNSIGNED", "NOT NULL, FK", "关联告警"],
            ["handler_user_id", "BIGINT UNSIGNED", "FK→sys_user.id", "当前处理人"],
            ["status", "TINYINT", "NOT NULL, DEFAULT 0", "0待处理/1处理中/2已解决/3已关闭"],
            ["title", "VARCHAR(200)", "NOT NULL", "工单标题"],
            ["description", "TEXT", "", "工单描述"],
            ["resolve_time", "DATETIME", "", "解决时间"],
        ],
    ),
    (
        "表3-6 ops_ticket_log（工单操作日志表）",
        ["字段名", "数据类型", "约束", "说明"],
        [
            ["id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "日志主键"],
            ["ticket_id", "BIGINT UNSIGNED", "NOT NULL, FK", "所属工单"],
            ["operator_id", "BIGINT UNSIGNED", "NOT NULL, FK", "操作人"],
            ["action", "VARCHAR(64)", "NOT NULL", "CREATE/START/RESOLVE/CLOSE等"],
            ["remark", "VARCHAR(500)", "", "操作备注"],
            ["operate_time", "DATETIME", "", "操作时间"],
        ],
    ),
    (
        "表3-7 ai_diagnosis（AI诊断报告表）",
        ["字段名", "数据类型", "约束", "说明"],
        [
            ["id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "报告主键"],
            ["alert_id", "BIGINT UNSIGNED", "NOT NULL, FK", "关联告警"],
            ["ticket_id", "BIGINT UNSIGNED", "FK", "关联工单"],
            ["ai_model", "VARCHAR(100)", "NOT NULL", "模型标识"],
            ["root_cause_analysis", "TEXT", "NOT NULL", "根因分析"],
            ["suggested_fix", "TEXT", "", "修复建议"],
            ["confidence_score", "DECIMAL(5,2)", "", "置信度0-100"],
            ["create_time", "DATETIME", "", "创建时间"],
        ],
    ),
    (
        "表3-8 ops_knowledge（知识库文档表）",
        ["字段名", "数据类型", "约束", "说明"],
        [
            ["id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "知识主键"],
            ["title", "VARCHAR(200)", "NOT NULL", "标题"],
            ["content_md", "MEDIUMTEXT", "NOT NULL", "Markdown正文"],
            ["sync_es_status", "TINYINT", "DEFAULT 0", "ES同步状态"],
            ["source_alert_id", "BIGINT UNSIGNED", "FK", "来源告警"],
            ["source_ticket_id", "BIGINT UNSIGNED", "FK", "来源工单"],
            ["source_diagnosis_id", "BIGINT UNSIGNED", "", "来源诊断"],
            ["lifecycle_status", "VARCHAR(32)", "", "生命周期状态"],
            ["entry_source", "VARCHAR(64)", "", "知识来源类型"],
        ],
    ),
    (
        "表3-9 ops_knowledge_audit_log（知识审计日志表）",
        ["字段名", "数据类型", "约束", "说明"],
        [
            ["id", "BIGINT UNSIGNED", "PK, AUTO_INCREMENT", "审计主键"],
            ["knowledge_id", "BIGINT UNSIGNED", "NOT NULL", "关联知识ID"],
            ["action", "VARCHAR(32)", "NOT NULL", "操作类型"],
            ["operator_id", "BIGINT UNSIGNED", "", "操作人ID"],
            ["from_status", "VARCHAR(32)", "", "变更前状态"],
            ["to_status", "VARCHAR(32)", "", "变更后状态"],
            ["operate_time", "DATETIME", "", "操作时间"],
        ],
    ),
]

MODULE_TABLES = [
    (
        "表3-10 告警接入模块核心类说明",
        ["类名", "层次", "职责"],
        [
            ["AlertWebhookController", "Controller", "接收 Alertmanager Webhook，解析 firing 告警"],
            ["AlertWorkflowServiceImpl", "Service", "告警入库、触发异步建单"],
            ["AlertTicketProducer", "MQ", "向 RabbitMQ 发送建单消息"],
            ["OpsAlert", "Entity", "告警实体，映射 ops_alert 表"],
        ],
    ),
    (
        "表3-11 异步建单模块 MQ 拓扑说明",
        ["组件", "名称", "说明"],
        [
            ["交换机", "ops.alert.exchange", "主业务直连交换机"],
            ["队列", "ops.alert.ticket.queue", "告警建单消费队列"],
            ["死信交换机", "ops.alert.dlx", "失败消息隔离"],
            ["死信队列", "ops.alert.ticket.dlq", "异常消息留存"],
            ["路由键", "alert.ticket.create", "建单消息路由"],
        ],
    ),
    (
        "表3-12 工单流转模块接口说明",
        ["接口", "方法", "说明"],
        [
            ["/ops-tickets/{id}/start", "POST", "开始处理，status 0→1"],
            ["/ops-tickets/{id}/assign", "PUT", "管理员指派处理人"],
            ["/ops-tickets/{id}/resolve", "POST", "标记已解决，status 1→2"],
            ["/ops-tickets/{id}/close", "POST", "关闭工单，status 2→3"],
        ],
    ),
]

MODULES_31 = [
    {
        "title": "（1）告警接入模块",
        "intro": (
            "告警接入模块负责接收 Prometheus 监控栈中 Alertmanager 推送的 Webhook 请求，"
            "将 firing 状态告警解析后写入 MySQL 的 ops_alert 表，并向消息队列发送建单通知。"
            "该模块是平台与外部监控系统对接的网关，要求快速响应、低耦合。"
        ),
        "impl": (
            "实现类为 AlertWebhookController 与 AlertWorkflowServiceImpl。"
            "Controller 层仅处理 status=firing 的告警项，从 labels 中提取 alertname、severity、instance，"
            "将 severity 映射为 1~4 等级后批量入库；入库成功后由 AlertTicketProducer 发送包含 alertId 的 JSON 消息。"
            "Webhook 接口 /api/alerts/webhook 在 WebMvcConfig 中排除 JWT 拦截，供 Alertmanager 直接调用。"
        ),
        "table_idx": 0,
        "activity": "【此处插入图：告警接入模块活动图】",
        "class_diag": "【此处插入图：告警接入模块类图】",
    },
    {
        "title": "（2）异步建单模块",
        "intro": (
            "异步建单模块基于 RabbitMQ 实现告警入库与工单创建的解耦。"
            "若 Webhook 同步完成建单和 AI 诊断，容易因大模型调用耗时而超时；"
            "因此采用“先 ACK 告警、后异步消费”的事件驱动模式。"
        ),
        "impl": (
            "RabbitMqConfig 声明交换机 ops.alert.exchange、队列 ops.alert.ticket.queue 及死信队列 ops.alert.ticket.dlq。"
            "AlertTicketConsumer 以手动 ACK 模式消费消息：成功则 basicAck，失败则 basicNack 进入死信队列。"
            "消费端调用 AlertWorkflowServiceImpl.createTicketFromAlert()，按 alert_id 幂等建单并触发默认模型诊断。"
        ),
        "table_idx": 1,
        "activity": "【此处插入图：异步建单模块活动图】",
        "class_diag": "【此处插入图：异步建单模块类图】",
    },
    {
        "title": "（3）工单流转模块",
        "intro": (
            "工单流转模块实现运维处置的标准化流程，将告警事件转化为可跟踪、可审计的工单对象。"
            "工单状态机包含四种状态：0-待处理、1-处理中、2-已解决、3-已关闭。"
        ),
        "impl": (
            "TicketKnowledgeServiceImpl 封装全部状态流转逻辑：startTicket 认领工单并设置处理人；"
            "assignHandler 仅管理员可改派；resolveTicket 填写经验总结并同步知识库；closeTicket 归档工单。"
            "每次操作写入 ops_ticket_log，并通过 assertOperatorCanProcess 校验处理人归属，防止越权。"
        ),
        "table_idx": 2,
        "activity": "【此处插入图：工单流转模块活动图】",
        "class_diag": "【此处插入图：工单流转模块类图】",
    },
    {
        "title": "（4）AI 诊断模块",
        "intro": (
            "AI 诊断模块在工单处置过程中调用大模型 API，根据告警上下文生成结构化的根因分析与修复建议报告。"
            "支持多模型配置，同一工单可按不同模型各生成一份报告，同工单同模型幂等。"
        ),
        "impl": (
            "AiDiagnosisWorkflowServiceImpl 负责 Prompt 构建、模型调用与结果解析入库；"
            "AiChatClientFactory 按模型供应商动态创建 Spring AI 客户端；"
            "AiModelServiceImpl 提供可选模型列表。告警建单后自动触发默认模型诊断，"
            "工单详情页支持手动选择模型再次诊断。"
        ),
        "table_idx": None,
        "activity": "【此处插入图：AI诊断模块活动图】",
        "class_diag": "【此处插入图：AI诊断模块类图】",
    },
    {
        "title": "（5）知识库检索模块",
        "intro": (
            "知识库模块负责故障处理经验的沉淀与检索，采用 MySQL 持久化 + Elasticsearch 全文检索的双存储架构。"
            "知识来源包括工单解决沉淀、诊断报告导入和手动录入三种。"
        ),
        "impl": (
            "TicketKnowledgeServiceImpl 在工单 resolve 时自动 upsert 知识条目并同步 ES 索引 ops_ticket_knowledge；"
            "KnowledgeArticleServiceImpl 支持草稿、审核、发布等生命周期管理；"
            "检索接口支持关键词多字段匹配与高亮展示，高亮失败时降级为普通查询。"
        ),
        "table_idx": None,
        "activity": "【此处插入图：知识库模块活动图】",
        "class_diag": "【此处插入图：知识库模块类图】",
    },
    {
        "title": "（6）认证授权模块",
        "intro": (
            "认证授权模块为全平台提供统一的身份认证与 RBAC 权限控制，"
            "定义 ADMIN、OPS、VIEWER 三种角色，分别对应系统管理、运维处置和只读查看权限。"
        ),
        "impl": (
            "AuthController 完成登录并签发 JWT；JwtAuthInterceptor 拦截业务接口校验 Token；"
            "SysUserServiceImpl 实现用户 CRUD、单用户单角色分配及内置 admin 账号保护。"
            "工单、知识等敏感操作在 Service 层再次校验角色与数据归属，形成前后端双重鉴权。"
        ),
        "table_idx": None,
        "activity": "【此处插入图：认证授权模块活动图】",
        "class_diag": "【此处插入图：认证授权模块类图】",
    },
]

MEMBER_B_SUMMARY = {
    "title": "组员B开发总结",
    "work": (
        "本人在项目中担任后端开发负责人，完成 Spring Boot 3 后端整体架构设计与实现。"
        "具体包括：Controller-Service-Mapper 三层结构搭建；9 张核心数据表设计与 MyBatis-Plus 实体映射；"
        "Alertmanager Webhook 告警接入；RabbitMQ 异步建单；四状态工单状态机与处理人权限校验；"
        "Spring AI 多模型诊断编排；知识库 MySQL+ES 双写检索；JWT+RBAC 认证授权；"
        "全部 REST API 设计与前后端联调。目前告警接入→异步建单→AI诊断→工单流转→知识沉淀的闭环已跑通。"
    ),
    "success": (
        "完成 12 个 Controller、20 余个 Service 实现类、9 张数据表建表脚本及 2 份迁移脚本；"
        "编写 mysql.sql、init-user.sql、inject-demo-alerts.ps1、start-backend.ps1 等工程脚本；"
        "实现告警 Webhook、工单状态机、MQ 死信队列、多模型 AI 诊断、ES 知识检索等核心能力；"
        "修复 admin 权限锁死、工单越权、知识库 404、ES 检索为空等多个联调问题。"
    ),
    "team": (
        "与组员 C 通过接口约定和 Vite 代理配置完成前后端联调；"
        "与组员 A 确定中间件 IP、Webhook 地址及 Docker 监控栈部署分工；"
        "与组员 D 共同确定 Prompt 结构与多模型 YAML 配置格式。"
        "后端坚持接口命名 RESTful、响应格式统一，减少前端适配成本。"
    ),
    "lessons": (
        "（1）配置与环境隔离：接手项目时需将硬编码 IP 改为团队虚拟机地址，AI 密钥通过 local.env 注入；"
        "（2）权限设计需三层：JWT 认证、RBAC 角色、业务数据归属校验；"
        "（3）异步解耦：RabbitMQ 手动 ACK+死信队列提升告警链路可靠性；"
        "（4）联调前检查网络连通性（MySQL/Redis/RabbitMQ/ES 四端口）；"
        "（5）ES 与 MySQL 双写应先保证基础检索可用再迭代高级功能。"
        "不足：ops_ticket.alert_id 缺少唯一索引、AI 与建单流程耦合偏紧、自动化测试覆盖不足。"
        "后续将引入多环境 Profile、拆分 AI 诊断独立队列、补充核心状态机单元测试。"
    ),
}


def fill_section_31(doc):
    # 3.1 节模块占位：逻辑架构段落后第一个 XXX模块（约段落 47）
    anchor = None
    for i, p in enumerate(doc.paragraphs):
        if i >= 38 and "模块" in p.text and "XXX" in p.text:
            anchor = p
            break
    if anchor is None and len(doc.paragraphs) > 47:
        anchor = doc.paragraphs[47]
    if anchor is None:
        raise RuntimeError("未找到 3.1 节 XXX模块 占位段落")

    # 清除占位：XXX/YYY/ZZZ 模块及"同样要求"
    to_remove = []
    started = False
    for p in doc.paragraphs:
        if p == anchor:
            started = True
        if started:
            if "物理架构" in p.text:
                break
            to_remove.append(p)

    # 保留第一个占位段，改写为总述
    set_paragraph_text(
        anchor,
        "在完成系统整体逻辑架构设计的基础上，本节按核心业务模块分别说明实现方案。"
        "各模块均给出文字描述、关键类/接口说明表，并预留 UML 活动图与类图（见插图）。"
        "系统整体逻辑架构如图3-1所示（请自行插入）。",
    )
    current = anchor

    for mod in MODULES_31:
        current = insert_paragraph_after(current, mod["title"])
        current = insert_paragraph_after(current, mod["intro"])
        current = insert_paragraph_after(current, mod["impl"])
        if mod["table_idx"] is not None:
            t = MODULE_TABLES[mod["table_idx"]]
            # 简化：只插标题行说明，完整表用 add_table
            pass
        current = insert_paragraph_after(current, mod["activity"])
        current = insert_paragraph_after(current, mod["class_diag"])
        current = insert_paragraph_after(current, "")

    # 删除旧占位段（除 anchor 外）
    for p in to_remove[1:]:
        if p._p.getparent() is not None:
            p._p.getparent().remove(p._p)

    # 在模块文字后补模块说明表（插在物理架构前）
    phys = find_paragraph(doc, "物理架构")
    if phys:
        prev = phys
        for cap, headers, rows in MODULE_TABLES:
            tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
            tbl.style = "Table Grid"
            for i, h in enumerate(headers):
                tbl.rows[0].cells[i].text = h
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    tbl.rows[ri + 1].cells[ci].text = val
            cap_p = insert_paragraph_after(prev, cap)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            prev._p.addnext(tbl._tbl)
            prev = cap_p


def fill_section_33(doc):
    anchor = None
    for i, p in enumerate(doc.paragraphs):
        if i >= 50 and "数据库" in p.text and "设计" in p.text:
            anchor = p
            break
    if anchor is None and len(doc.paragraphs) > 55:
        anchor = doc.paragraphs[55]
    if anchor is None:
        raise RuntimeError("未找到数据库设计段落")

    set_paragraph_text(
        anchor,
        "数据库设计",
    )
    intro = insert_paragraph_after(
        anchor,
        "本系统业务主数据存储于 MySQL 8.0 数据库 ops_ai_platform，知识检索索引存储于 Elasticsearch。"
        "共设计 9 张核心表，满足用户权限、告警、工单、诊断、知识库等业务闭环需求。"
        "数据库 E-R 关系如图3-5所示（请自行插入）。设计原则：外键保证关联一致性；"
        "ops_alert.raw_payload 保留原始 JSON 供 AI 诊断；知识库采用 MySQL+ES 双写。",
    )

    current = intro
    for cap, headers, rows in DB_TABLES:
        cap_p = insert_paragraph_after(current, cap)
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap_p.runs:
            run.bold = True
        tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = "Table Grid"
        for i, h in enumerate(headers):
            tbl.rows[0].cells[i].text = h
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                tbl.rows[ri + 1].cells[ci].text = val
        current._p.addnext(tbl._tbl)
        current = cap_p
        current = insert_paragraph_after(current, "")


def fill_member_summary(doc):
    # 组员 B → 模板中第 2 个成员总结占位（YYY开发总结，约段落 95）
    anchors = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.endswith("开发总结") and len(t) < 20:
            anchors.append((i, p))
    if len(anchors) < 2:
        # 回退：按固定索引
        if len(doc.paragraphs) > 95:
            anchor = doc.paragraphs[95]
        else:
            raise RuntimeError(f"未找到成员开发总结占位，仅找到 {len(anchors)} 个")
    else:
        anchor = anchors[1][1]

    set_paragraph_text(anchor, MEMBER_B_SUMMARY["title"])

    sections = [
        ("承担工作", MEMBER_B_SUMMARY["work"]),
        ("工作成果", MEMBER_B_SUMMARY["success"]),
        ("团队协调", MEMBER_B_SUMMARY["team"]),
        ("收获与教训", MEMBER_B_SUMMARY["lessons"]),
    ]

    current = anchor
    for heading, body in sections:
        current = insert_paragraph_after(current, heading)
        current = insert_paragraph_after(current, body)

    # 删除旧占位"要求：..."段落
    for p in list(doc.paragraphs):
        if p.text.strip().startswith("要求：") and "组员" not in p.text:
            if p._p.getparent() is not None:
                p._p.getparent().remove(p._p)


def main():
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)
    shutil.copy2(TEMPLATE, OUTPUT)
    doc = Document(str(OUTPUT))
    fill_section_31(doc)
    fill_section_33(doc)
    fill_member_summary(doc)
    doc.save(str(OUTPUT))
    print(f"已生成: {OUTPUT}")


if __name__ == "__main__":
    main()
